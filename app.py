import io
import json
import os
import re
from typing import Any, Dict, List

import streamlit as st
from google import genai
from google.genai import types
from pypdf import PdfReader
from docx import Document


st.set_page_config(
    page_title="Resume ATS Analyzer",
    page_icon="📄",
    layout="wide",
)

st.title("📄 Resume ATS Analyzer")
st.caption("Upload a resume, get an ATS-style score, and receive actionable improvements.")

# -----------------------------
# Configuration
# -----------------------------
MODEL_NAME = os.getenv("GEMINI_MODEL", "gemini-3.6-flash")


def get_api_key() -> str:
    """Read Gemini API key from Streamlit secrets first, then environment."""
    try:
        key = st.secrets.get("GEMINI_API_KEY")
    except Exception:
        key = None

    return key or os.getenv("GEMINI_API_KEY", "")


def extract_pdf_text(file_bytes: bytes) -> str:
    reader = PdfReader(io.BytesIO(file_bytes))
    pages = []
    for page in reader.pages:
        pages.append(page.extract_text() or "")
    return "\n".join(pages).strip()


def extract_docx_text(file_bytes: bytes) -> str:
    document = Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]

    # Also capture simple table content, since some resumes use tables.
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                paragraphs.append(" | ".join(cells))

    return "\n".join(paragraphs).strip()


def extract_text(uploaded_file) -> str:
    data = uploaded_file.getvalue()
    suffix = uploaded_file.name.lower().rsplit(".", 1)[-1]

    if suffix == "pdf":
        return extract_pdf_text(data)
    if suffix == "docx":
        return extract_docx_text(data)

    raise ValueError("Unsupported file type. Please upload a PDF or DOCX resume.")


def clean_json_text(text: str) -> str:
    """Remove accidental Markdown code fences around JSON."""
    text = text.strip()
    text = re.sub(r"^```(?:json)?\s*", "", text, flags=re.IGNORECASE)
    text = re.sub(r"\s*```$", "", text)
    return text.strip()


def analyze_resume(resume_text: str, job_description: str, api_key: str) -> Dict[str, Any]:
    client = genai.Client(api_key=api_key)

    prompt = f"""
You are an expert ATS (Applicant Tracking System) resume evaluator and professional resume coach.

Analyze the resume below. If a job description is provided, evaluate the resume against it.
Do NOT invent experience, education, certifications, metrics, or skills for the candidate.

Return ONLY valid JSON matching the schema below. No Markdown, no code fences, and no extra text.

Scoring rules:
- Overall ATS score: 0-100.
- If a job description is provided, weight keyword/job-match relevance heavily.
- If no job description is provided, score general ATS readiness.
- Consider: contact information, section headings, readability, standard formatting,
  keyword relevance, measurable achievements, skills, experience clarity, education,
  chronology, and likely parsing problems.
- Give concise, practical recommendations.

JSON schema:
{{
  "overall_score": 0,
  "score_breakdown": {{
    "formatting_and_parsing": 0,
    "keyword_optimization": 0,
    "experience_and_achievements": 0,
    "skills": 0,
    "structure_and_clarity": 0
  }},
  "summary": "short overall assessment",
  "strengths": ["strength 1", "strength 2"],
  "improvements": [
    {{
      "priority": "High",
      "issue": "specific issue",
      "why_it_matters": "why it matters for ATS/recruiters",
      "recommendation": "specific action"
    }}
  ],
  "keywords_to_add": ["keyword 1", "keyword 2"],
  "formatting_warnings": ["warning 1"],
  "missing_sections_or_information": ["item 1"],
  "ats_safe_format_tips": ["tip 1", "tip 2"]
}}

Keep score_breakdown values between 0 and 100.
If something is not applicable, use a reasonable score and explain it briefly.
Limit improvements, keywords, warnings, and tips to the most useful items.

RESUME:
---BEGIN RESUME---
{resume_text[:30000]}
---END RESUME---

JOB DESCRIPTION:
---BEGIN JOB DESCRIPTION---
{job_description[:20000] if job_description.strip() else "No job description provided. Evaluate general ATS readiness."}
---END JOB DESCRIPTION---
"""

    response = client.models.generate_content(
        model=MODEL_NAME,
        contents=prompt,
        config=types.GenerateContentConfig(
            temperature=0.2,
            response_mime_type="application/json",
        ),
    )

    raw = getattr(response, "text", "") or ""
    if not raw.strip():
        raise RuntimeError("Gemini returned an empty response.")

    try:
        return json.loads(clean_json_text(raw))
    except json.JSONDecodeError as exc:
        raise RuntimeError("Gemini returned invalid JSON. Please try again.") from exc


def safe_score(value: Any) -> int:
    try:
        return max(0, min(100, int(round(float(value)))))
    except (TypeError, ValueError):
        return 0


# -----------------------------
# UI
# -----------------------------
with st.sidebar:
    st.header("Settings")
    st.info(
        "Add GEMINI_API_KEY to Streamlit secrets or your environment before analyzing."
    )
    st.caption(f"Model: {MODEL_NAME}")

uploaded_file = st.file_uploader(
    "Upload your resume",
    type=["pdf", "docx"],
    help="PDF or DOCX. Text-based PDFs work best.",
)

job_description = st.text_area(
    "Optional: paste the job description",
    height=220,
    placeholder="Paste the job description here to get a more job-specific ATS score...",
)

analyze_button = st.button(
    "🚀 Analyze Resume",
    type="primary",
    use_container_width=True,
)

if analyze_button:
    if uploaded_file is None:
        st.error("Please upload a PDF or DOCX resume first.")
        st.stop()

    api_key = get_api_key()
    if not api_key:
        st.error("Gemini API key is missing. Add GEMINI_API_KEY to Streamlit secrets.")
        st.stop()

    try:
        with st.spinner("Extracting resume text..."):
            resume_text = extract_text(uploaded_file)

        if not resume_text.strip():
            st.error(
                "No readable text was extracted. If this is a scanned/image-only PDF, "
                "please upload a text-based PDF or DOCX."
            )
            st.stop()

        with st.spinner("Analyzing your resume with Gemini..."):
            result = analyze_resume(resume_text, job_description, api_key)

        score = safe_score(result.get("overall_score"))

        st.session_state["analysis"] = result
        st.session_state["score"] = score

    except Exception as exc:
        st.error(f"Analysis failed: {exc}")
        st.stop()


if "analysis" in st.session_state:
    result = st.session_state["analysis"]
    score = st.session_state["score"]

    st.divider()

    col1, col2 = st.columns([1, 2])

    with col1:
        st.metric("ATS Score", f"{score}/100")

    with col2:
        st.progress(score / 100)
        st.write(result.get("summary", "No summary returned."))

    st.subheader("📊 Score Breakdown")

    breakdown = result.get("score_breakdown", {})
    cols = st.columns(5)
    labels = [
        ("Formatting & Parsing", "formatting_and_parsing"),
        ("Keywords", "keyword_optimization"),
        ("Experience", "experience_and_achievements"),
        ("Skills", "skills"),
        ("Structure", "structure_and_clarity"),
    ]

    for col, (label, key) in zip(cols, labels):
        with col:
            st.metric(label, f"{safe_score(breakdown.get(key))}/100")

    left, right = st.columns(2)

    with left:
        st.subheader("✅ Strengths")
        strengths = result.get("strengths", [])
        if strengths:
            for item in strengths:
                st.markdown(f"- {item}")
        else:
            st.write("No strengths returned.")

    with right:
        st.subheader("🔑 Keywords to Add")
        keywords = result.get("keywords_to_add", [])
        if keywords:
            st.write(", ".join(str(k) for k in keywords))
        else:
            st.write("No additional keywords identified.")

    st.subheader("🛠️ Recommended Improvements")
    improvements: List[Dict[str, Any]] = result.get("improvements", [])

    if improvements:
        for i, item in enumerate(improvements, start=1):
            priority = item.get("priority", "Medium")
            issue = item.get("issue", "Improvement")
            with st.expander(f"{i}. [{priority}] {issue}"):
                st.write(f"**Why it matters:** {item.get('why_it_matters', '')}")
                st.write(f"**What to do:** {item.get('recommendation', '')}")
    else:
        st.write("No specific improvements returned.")

    col_a, col_b = st.columns(2)

    with col_a:
        st.subheader("⚠️ Formatting Warnings")
        for item in result.get("formatting_warnings", []) or ["No major formatting warnings."]:
            st.markdown(f"- {item}")

    with col_b:
        st.subheader("📌 Missing Sections / Information")
        for item in result.get("missing_sections_or_information", []) or ["None identified."]:
            st.markdown(f"- {item}")

    st.subheader("🧾 ATS-Safe Formatting Tips")
    for item in result.get("ats_safe_format_tips", []) or ["Use simple, standard resume formatting."]:
        st.markdown(f"- {item}")

    # Let the user save/share the structured result.
    st.download_button(
        "⬇️ Download Analysis JSON",
        data=json.dumps(result, indent=2, ensure_ascii=False),
        file_name="resume_ats_analysis.json",
        mime="application/json",
    )

st.divider()
st.caption(
    "Note: This is an ATS-style estimate, not a score produced by a specific employer's ATS. "
    "Different ATS platforms and job postings use different ranking rules."
)
